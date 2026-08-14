from __future__ import annotations

import asyncio
import json
import sqlite3
from pathlib import Path

import pytest

import backend.rag.evaluate as rag_evaluation
from docx import Document
from pypdf import PdfWriter
from pydantic import ValidationError

from backend.agent import LocalAgent, _contextualize_rag_query, _evidence_query_anchors, _paragraphs_are_cited, _parse_rag_answer, _strip_thinking, _thinking
from backend.contracts import ChatRequest
from backend.modes import ChatModes
from backend.rag.chunking import chunk_text
from backend.rag.citations import build_citations, validate_citations
from backend.rag.documents import GotOcrVramSwap, convert_to_markdown
from backend.rag.embeddings import EmbeddingClient, index_embeddings
from backend.rag.ingest import ingest_documents, reindex_all
from backend.rag.search import hybrid_search, reciprocal_rank_fusion, vector_search
from backend.rag.rerank import rerank
from backend.rag.store import RagStore
from backend.web import parse_search_results


@pytest.fixture()
def corpus(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    workspace = tmp_path / "workspace"
    docs = workspace / "docs"
    docs.mkdir(parents=True)
    store = RagStore(workspace / ".rag_index")
    monkeypatch.setattr("backend.rag.ingest.WORKSPACE_ROOT", workspace)
    monkeypatch.setattr("backend.rag.ingest.RAG_DOCS_DIR", docs)
    monkeypatch.setattr("backend.rag.ingest.RAG_SOURCES_DIR", workspace / ".rag_sources")
    monkeypatch.setattr("backend.rag.documents.WORKSPACE_ROOT", workspace)
    monkeypatch.setattr("backend.rag.documents.RAG_SOURCES_DIR", workspace / ".rag_sources")
    monkeypatch.setattr("backend.rag.documents.RAG_CACHE_DIR", workspace / ".rag_cache")
    return workspace, docs, store


def test_contract_defaults_and_validation():
    legacy = ChatRequest(message="hola")
    assert legacy.modes is None
    modern = ChatRequest(message="hola", modes={"rag": True})
    assert modern.modes == ChatModes(chat=True, rag=True)
    with pytest.raises(ValidationError):
        ChatRequest(message="")


def test_unclosed_thinking_is_never_exposed_as_answer():
    raw = "<think>analisis con una cita [1] sin terminar"
    assert _strip_thinking(raw) == ""
    assert _thinking(raw) == "analisis con una cita [1] sin terminar"
    assert _paragraphs_are_cited("Hecho [1].\n\nInferencia: posible conclusion.")
    assert not _paragraphs_are_cited("Hecho [1].\n\nOtro hecho sin cita.")
    assert not _paragraphs_are_cited("[1][2][3][4]")


def test_structured_rag_answer_requires_text_and_valid_source_ids():
    raw = json.dumps(
        {
            "paragraphs": [
                {"text": "La Union establece normas armonizadas.", "source_ids": [2, 1, 2]},
                {"text": "[3]", "source_ids": [3]},
                {"text": "Fuente inexistente", "source_ids": [9]},
            ]
        }
    )
    answer, referenced = _parse_rag_answer(raw, 3)
    assert answer == "La Union establece normas armonizadas. [1][2]"
    assert referenced == {1, 2}
    assert _parse_rag_answer('{"paragraphs": []}', 3) == ("", set())


def test_rag_followup_uses_previous_user_question_as_search_context():
    history = [
        {"role": "user", "content": "¿Qué regula la UE sobre inteligencia artificial?"},
        {"role": "assistant", "content": "Existe un reglamento europeo [1]."},
    ]
    query = _contextualize_rag_query("¿Y qué sanciones contempla?", history)
    assert "Qué regula la UE" in query and "sanciones" in query
    assert _contextualize_rag_query("Explica la fotosintesis con detalle", history) == "Explica la fotosintesis con detalle"


def test_evidence_requires_a_distinctive_query_anchor():
    chunks = [{"content": "El Reglamento europeo establece una politica armonizada para sistemas de IA."}]
    assert _evidence_query_anchors("Que politica existe en la UE sobre IA?", chunks) == {"politica", "ia"}
    assert not _evidence_query_anchors("Que novedades incluye Python 3.14?", chunks)


def test_duckduckgo_results_are_parsed_into_public_web_evidence():
    document = """
    <a class="result__a" href="//duckduckgo.com/l/?uddg=https%3A%2F%2Fexample.com%2Fai">AI Act</a>
    <a class="result__snippet">Resumen verificable del reglamento europeo.</a>
    <a class="result__a" href="http://127.0.0.1/private">Privado</a>
    <a class="result__snippet">No debe exponerse.</a>
    """
    assert parse_search_results(document) == [
        {"url": "https://example.com/ai", "title": "AI Act", "snippet": "Resumen verificable del reglamento europeo."}
    ]


def test_rag_falls_back_to_cited_web_results(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setattr("backend.agent.hybrid_search", lambda *args, **kwargs: ([], {"selected": []}))

    async def fake_search(client, query):
        return [{"url": "https://example.com/ai", "title": "AI Act", "snippet": "La UE aplica normas comunes a la IA."}]

    async def fake_complete(messages, **kwargs):
        return json.dumps({"paragraphs": [{"text": "La UE aplica normas comunes a la IA.", "source_ids": [1]}]})

    monkeypatch.setattr("backend.agent.search_web", fake_search)
    agent = LocalAgent()
    monkeypatch.setattr(agent, "_complete_text", fake_complete)
    result = asyncio.run(agent.chat("politica europea de IA", [], ChatModes(chat=True, rag=True, web=True)))
    asyncio.run(agent.close())
    assert result["stopped"] == "final" and result["citations"][0]["source_type"] == "web"
    assert result["citations"][0]["path"] == "https://example.com/ai"


def test_rag_uses_web_when_local_chunks_do_not_answer(monkeypatch: pytest.MonkeyPatch):
    irrelevant = {
        "id": "local-1",
        "path": "docs/ia.md",
        "title": "IA",
        "section": "General",
        "start_line": 1,
        "end_line": 2,
        "content": "Normas europeas sobre inteligencia artificial.",
        "score": 0.1,
        "reasons": ["keyword"],
    }
    monkeypatch.setattr("backend.agent.hybrid_search", lambda *args, **kwargs: ([irrelevant], {"selected": ["local-1"]}))

    async def fake_search(client, query):
        return [{"url": "https://example.com/python", "title": "Python 3.14", "snippet": "Python 3.14 incorpora nuevas funciones."}]

    responses = iter([json.dumps({"paragraphs": [{"text": "Python 3.14 incorpora nuevas funciones.", "source_ids": [1]}]})])

    async def fake_complete(messages, **kwargs):
        return next(responses)

    monkeypatch.setattr("backend.agent.search_web", fake_search)
    agent = LocalAgent()
    monkeypatch.setattr(agent, "_complete_text", fake_complete)
    result = asyncio.run(agent.chat("Novedades de Python 3.14", [], ChatModes(chat=True, rag=True, web=True)))
    asyncio.run(agent.close())
    assert result["stopped"] == "final" and result["citations"][0]["source_type"] == "web"


def test_chunking_preserves_sections_symbols_and_lines(tmp_path: Path):
    markdown = "# Uno\n\ncontenido uno\n## Dos\n\ncontenido dos"
    chunks = chunk_text(tmp_path / "a.md", markdown, 10, 2)
    assert chunks[0].start_line == 1
    assert chunks[-1].end_line == 6
    assert any(chunk.section == "Uno > Dos" for chunk in chunks)
    code = "def first():\n    return 1\n\nclass Second:\n    pass\n"
    symbols = {chunk.symbol for chunk in chunk_text(tmp_path / "a.py", code, 50, 0)}
    assert {"first", "Second"} <= symbols
    javascript = "export function alpha() { return 1; }\nconst beta = () => 2;\n"
    js_symbols = {chunk.symbol for chunk in chunk_text(tmp_path / "a.js", javascript, 50, 0)}
    assert {"alpha", "beta"} <= js_symbols


def test_incremental_ingestion_fts_and_orphan_cleanup(corpus):
    _, docs, store = corpus
    source = docs / "manual.md"
    source.write_text("# GPU\n\nUsa LLAMA_PARALLEL=1 para ahorrar memoria.\n", encoding="utf-8")
    first = ingest_documents(docs, store)
    second = ingest_documents(docs, store)
    assert first["indexed"] == 1 and first["chunks"] >= 1
    assert second["unchanged"] == 1 and second["indexed"] == 0
    hits = store.keyword_search("LLAMA_PARALLEL")
    assert hits and hits[0]["path"] == "docs/manual.md"
    source.write_text("# GPU\n\nEl contexto recomendado es 65536.\n", encoding="utf-8")
    changed = ingest_documents(docs, store)
    assert changed["indexed"] == 1
    assert not store.keyword_search("LLAMA_PARALLEL")
    source.unlink()
    deleted = ingest_documents(docs, store)
    assert deleted["deleted"] == 1 and store.status()["chunks"] == 0


def test_cascade_deleting_a_document_leaves_no_orphaned_fts_rows(corpus):
    _, docs, store = corpus
    (docs / "manual.md").write_text("# GPU\n\nUsa LLAMA_PARALLEL=1 para ahorrar memoria.\n", encoding="utf-8")
    ingest_documents(docs, store)
    with store.connect() as db:
        doc_id = db.execute("SELECT id FROM documents WHERE path='docs/manual.md'").fetchone()[0]
        assert db.execute("SELECT count(*) FROM chunk_fts").fetchone()[0] > 0
        # Simulates any caller that deletes a document directly and relies on
        # `chunks.document_id ... ON DELETE CASCADE` rather than going through
        # ingest_documents' own chunk_fts cleanup (e.g. a future delete-document tool,
        # or manual DB maintenance) -- FK cascade reaches `chunks`, but FTS5 virtual
        # tables aren't part of SQLite's FK graph, so only the chunk_fts_ad trigger
        # keeps this consistent.
        db.execute("DELETE FROM documents WHERE id=?", (doc_id,))
    with store.connect() as db:
        assert db.execute("SELECT count(*) FROM chunk_fts").fetchone()[0] == 0


def test_full_rebuild_includes_docs_and_canonical_sources(corpus):
    workspace, docs, store = corpus
    (docs / "manual.md").write_text("contenido original", encoding="utf-8")
    sources = workspace / ".rag_sources"
    sources.mkdir()
    (sources / "converted.md").write_text("contenido convertido", encoding="utf-8")
    result = reindex_all(store)
    assert result["indexed"] == 2
    assert store.status()["documents"] == 2
    assert store.keyword_search("convertido")


def test_embeddings_incremental_model_invalidation_and_vector_search(corpus):
    _, docs, store = corpus
    (docs / "animals.md").write_text("# Felinos\n\nEl gato duerme sobre el teclado.\n", encoding="utf-8")
    (docs / "space.md").write_text("# Espacio\n\nLa nave explora una galaxia distante.\n", encoding="utf-8")
    ingest_documents(docs, store)
    client = EmbeddingClient(provider="deterministic", model="test-v1", dimensions=64)
    first = index_embeddings(store, client)
    second = index_embeddings(store, client)
    changed_model = index_embeddings(store, EmbeddingClient(provider="deterministic", model="test-v2", dimensions=64))
    assert first["indexed"] == 2 and second["indexed"] == 0 and changed_model["indexed"] == 2
    hits = vector_search("galaxia", store, client)
    assert hits[0]["path"] == "docs/space.md"


def test_rrf_is_stable_and_hybrid_has_trace(corpus):
    _, docs, store = corpus
    (docs / "one.md").write_text("# Config\n\nalpha beta comando exacto\n", encoding="utf-8")
    ingest_documents(docs, store)
    rows = store.keyword_search("comando")
    fused = reciprocal_rank_fusion([rows, rows])
    assert fused[0]["id"] == rows[0]["id"] and fused[0]["score"] > 0
    selected, trace = hybrid_search("comando exacto", store, top_k=3)
    assert selected and trace["keyword_ranking"] and trace["selected"]


def test_low_confidence_vector_only_result_is_rejected(corpus):
    _, docs, store = corpus
    (docs / "only.md").write_text("documentacion sobre volcanes", encoding="utf-8")
    ingest_documents(docs, store)
    client = EmbeddingClient(provider="deterministic", model="test", dimensions=32)
    index_embeddings(store, client)
    selected, trace = hybrid_search("contabilidad bancaria", store, client=client)
    assert not selected and not trace["keyword_ranking"]


def test_stopwords_do_not_create_false_lexical_evidence(corpus):
    _, docs, store = corpus
    (docs / "gpu.md").write_text("El modelo se ejecuta en la GPU.", encoding="utf-8")
    ingest_documents(docs, store)
    assert not store.keyword_search("Cual es la capital de Marte")


def test_hybrid_filters_tags_and_dates(corpus):
    _, docs, store = corpus
    (docs / "tagged.md").write_text("---\ntags: [gpu, local]\n---\n# GPU\n\nVulkan activo\n", encoding="utf-8")
    ingest_documents(docs, store)
    selected, _ = hybrid_search("Vulkan", store, top_k=3, filters={"tag": "gpu", "type": "md", "path": "docs/"})
    assert selected
    missing, _ = hybrid_search("Vulkan", store, top_k=3, filters={"tag": "cloud"})
    assert not missing


def test_citations_reference_existing_chunks(corpus):
    _, docs, store = corpus
    (docs / "source.txt").write_text("Una fuente verificable.\n", encoding="utf-8")
    ingest_documents(docs, store)
    chunks = store.keyword_search("verificable")
    citations = build_citations(chunks)
    assert validate_citations(citations, store)
    assert citations[0].start_line == 1 and citations[0].path == "docs/source.txt"


def test_llm_rerank_is_strict_cached_and_keeps_only_candidate_ids(corpus):
    _, _, store = corpus
    candidates = [{"id": "a"}, {"id": "b"}, {"id": "c"}]
    calls = []

    async def completion(prompt):
        calls.append(prompt)
        return json.dumps({"ranked_chunk_ids": ["c", "unknown", "a"]})

    ranked, trace = asyncio.run(rerank("query", candidates, completion, store, top_k=2))
    cached, cache_trace = asyncio.run(rerank("query", candidates, completion, store, top_k=2))
    assert [item["id"] for item in ranked] == ["c", "a"]
    assert [item["id"] for item in cached] == ["c", "a"]
    assert trace["strategy"] == "llm" and cache_trace["strategy"] == "llm_cache" and len(calls) == 1


def test_rerank_cannot_demote_stronger_hybrid_evidence(corpus):
    _, _, store = corpus
    candidates = [
        {"id": "strong", "content": "exact", "reasons": ["keyword", "vector"]},
        {"id": "weak", "content": "semantic", "reasons": ["vector"]},
        {"id": "other", "content": "semantic", "reasons": ["vector"]},
    ]

    async def completion(prompt):
        return json.dumps({"ranked_chunk_ids": ["weak", "other", "strong"]})

    ranked, _ = asyncio.run(rerank("query", candidates, completion, store, top_k=2))
    assert ranked[0]["id"] == "strong"


def test_docx_and_image_ocr_become_canonical_markdown(corpus, monkeypatch: pytest.MonkeyPatch):
    workspace, _, store = corpus
    document = Document()
    document.add_heading("Instalacion", level=1)
    document.add_paragraph("Ejecuta el comando local.")
    docx_path = workspace / "guide.docx"
    document.save(docx_path)
    result = convert_to_markdown(docx_path, store)
    assert "# Instalacion" in (workspace / result["markdown"]).read_text(encoding="utf-8")
    image = workspace / "scan.png"
    image.write_bytes(b"fake-image")
    monkeypatch.setattr("backend.rag.documents._ocr_image", lambda path, provider="tesseract": "Texto reconocido de pagina")
    scan = convert_to_markdown(image, store)
    assert "Texto reconocido" in (workspace / scan["markdown"]).read_text(encoding="utf-8")
    assert store.status()["sources"] == 2
    assert store.keyword_search("reconocido")


def test_scanned_pdf_runs_page_ocr_and_is_indexed(corpus, monkeypatch: pytest.MonkeyPatch):
    workspace, _, store = corpus
    pdf = workspace / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=200)
    with pdf.open("wb") as output:
        writer.write(output)
    monkeypatch.setattr("backend.rag.documents._ocr_image", lambda path, provider="rapidocr": "Factura escaneada numero 31415")
    result = convert_to_markdown(pdf, store)
    assert result["ocr_pages"] == [1]
    assert store.keyword_search("31415")[0]["path"].startswith(".rag_sources/")


def test_chat_only_never_uses_tool(monkeypatch: pytest.MonkeyPatch):
    agent = LocalAgent()
    called = []

    async def complete(messages, max_tokens=512, **kwargs):
        called.append(messages)
        return "Respuesta normal"

    monkeypatch.setattr(agent, "_complete_text", complete)
    result = asyncio.run(agent.chat("lista archivos", [], ChatModes(chat=True)))
    asyncio.run(agent.close())
    assert result["answer"] == "Respuesta normal"
    assert all(event.get("type") != "tool" for event in result["trace"])
    assert len(called) == 1


def test_frontend_persists_modes_and_renders_trace():
    script = Path("frontend/app.js").read_text(encoding="utf-8")
    html = Path("frontend/index.html").read_text(encoding="utf-8")
    assert "localStorage.setItem" in script and "modes" in script
    assert "HISTORY_KEY" in script and "saveHistory" in script and "localStorage.removeItem(HISTORY_KEY)" in script
    assert "history.slice(0, -1)" in script and "MAX_HISTORY_MESSAGES = 40" in script
    assert "LOADING_PHRASES" in script and "setInterval" in script and 'citation.source_type === "web"' in script
    assert "copy-trace" in html and 'id="mode-button"' in html and 'id="clear-chat"' in html



def test_plain_chat_retries_incomplete_thinking_and_keeps_history(monkeypatch: pytest.MonkeyPatch):
    agent = LocalAgent()
    calls = []
    responses = iter([
        "<think>razonamiento sin cerrar",
        '<think>breve</think>{"final":"AZULEJO-731"}',
    ])

    async def complete(messages, **kwargs):
        calls.append((messages, kwargs))
        return next(responses)

    monkeypatch.setattr(agent, "_complete_text", complete)
    result = asyncio.run(
        agent.chat(
            "¿Cuál es mi código?",
            [{"role": "user", "content": "Mi código es AZULEJO-731."}],
            ChatModes(chat=True, reasoning_panel=True),
        )
    )
    asyncio.run(agent.close())
    assert result["answer"] == "AZULEJO-731"
    assert result["stopped"] == "final"
    assert len(calls) == 2
    assert any("AZULEJO-731" in item["content"] for item in calls[0][0])
    assert calls[0][1]["max_tokens"] == -1
    assert calls[1][1]["max_tokens"] == -1
    assert "response_format" not in calls[0][1]


def test_got_ocr_vram_swap_unloads_both_and_restores_both(monkeypatch: pytest.MonkeyPatch):
    events = []
    monkeypatch.setattr("backend.rag.documents._router_request", lambda action, model: events.append((action, model)))
    with GotOcrVramSwap():
        assert set(events) == {("unload", "local-gguf"), ("unload", "bge-small-en-v1.5")}

    assert events[-2:] == [("load", "bge-small-en-v1.5"), ("load", "local-gguf")]


def test_got_ocr_vram_swap_restores_both_after_failure(monkeypatch: pytest.MonkeyPatch):
    events = []
    monkeypatch.setattr("backend.rag.documents._router_request", lambda action, model: events.append((action, model)))
    with pytest.raises(RuntimeError, match="ocr failed"):
        with GotOcrVramSwap():
            raise RuntimeError("ocr failed")
    assert events[-2:] == [("load", "bge-small-en-v1.5"), ("load", "local-gguf")]


def test_router_enabled_conversion_without_ocr_never_touches_router(corpus, monkeypatch: pytest.MonkeyPatch):
    workspace, _, store = corpus
    events = []
    monkeypatch.setattr("backend.rag.documents._router_request", lambda action, model: events.append((action, model)))
    monkeypatch.setattr("backend.rag.documents.MODEL_ROUTER_ENABLED", True)
    monkeypatch.setattr("backend.rag.embeddings.index_embeddings", lambda store=None: {"indexed": 0})

    document = Document()
    document.add_heading("Instalacion", level=1)
    document.add_paragraph("Ejecuta el comando local.")
    docx_path = workspace / "guide.docx"
    document.save(docx_path)
    result = convert_to_markdown(docx_path, store)

    assert events == []
    assert result["embedding_index"] == {"indexed": 0}



def test_evaluation_scores_generated_citations_not_only_retrieval(monkeypatch: pytest.MonkeyPatch):
    retrieved = [{"path": "docs/runtime.md", "content": "valor 1"}]
    monkeypatch.setattr(rag_evaluation, "hybrid_search", lambda *args, **kwargs: (retrieved, {"selected": ["chunk"]}))

    class FakeAgent:
        def __init__(self, **kwargs):
            pass

        async def chat(self, *args, **kwargs):
            return {
                "answer": "El valor es 1 [1].",
                "citations": [{"id": 1, "path": "docs/incorrecto.md"}],
                "stopped": "final",
            }

        async def close(self):
            pass

    monkeypatch.setattr(rag_evaluation, "LocalAgent", FakeAgent)
    results = asyncio.run(
        rag_evaluation._evaluate_cases(
            [{"question": "valor", "expected_answer": "1", "expected_sources": ["docs/runtime.md"]}],
            object(),
            None,
            5,
        )
    )
    assert results[0]["hit"] is True
    assert results[0]["citation_correct"] is False
    assert results[0]["grounded"] is False
