from __future__ import annotations

import hashlib
import json
import shutil
import subprocess
import sys
import tempfile
import time
from contextlib import AbstractContextManager
from functools import wraps
from pathlib import Path
from typing import Any

from backend.config import EMBEDDINGS_MODEL, LLM_MODEL, MODEL_ROUTER_BASE_URL, MODEL_ROUTER_ENABLED, MODEL_ROUTER_RESTORE_CHAT, OCR_ENABLED, OCR_MODEL, OCR_MODEL_DIR, OCR_PROVIDER, RAG_CACHE_DIR, RAG_SOURCES_DIR, WORKSPACE_ROOT
from backend.rag.store import RagStore


def _safe_input(path: Path) -> Path:
    path = path.resolve()
    try:
        path.relative_to(WORKSPACE_ROOT)
    except ValueError as exc:
        raise ValueError("Document path escapes workspace") from exc
    if not path.is_file():
        raise FileNotFoundError(path)
    return path


def _cache_key(data: bytes, provider: str) -> str:
    return hashlib.sha256(data + provider.encode()).hexdigest()


def _model_status(model: str) -> str | None:
    import httpx

    response = httpx.get(f"{MODEL_ROUTER_BASE_URL.rstrip('/')}/v1/models", timeout=30)
    response.raise_for_status()
    for entry in response.json().get("data", []):
        if entry.get("id") == model:
            return entry.get("status", {}).get("value")
    return None


def _router_request(action: str, model: str, ready_timeout: float = 120.0) -> None:
    import httpx

    response = httpx.post(
        f"{MODEL_ROUTER_BASE_URL.rstrip('/')}/models/{action}",
        json={"model": model},
        timeout=300,
    )
    if response.status_code == 400 and action == "unload" and "not running" in response.text:
        return
    response.raise_for_status()
    if response.json().get("success") is not True:
        raise RuntimeError(f"Router could not {action} model {model}")

    # /models/{action} reports success as soon as the request is accepted, not once the
    # model has actually finished loading/unloading (it stays "loading" for a bit first).
    # Callers issue requests against `model` right after this returns, so without this
    # poll they race a still-"loading" model and get a 400 "model is not loaded".
    expected = "loaded" if action == "load" else "unloaded"
    deadline = time.monotonic() + ready_timeout
    status = None
    while time.monotonic() < deadline:
        status = _model_status(model)
        if status in (expected, None):
            return
        time.sleep(0.3)
    raise RuntimeError(f"Router {action} of {model} did not reach '{expected}' within {ready_timeout}s (last status: {status})")


class GotOcrVramSwap(AbstractContextManager["GotOcrVramSwap"]):
    """Frees VRAM for the GOT-OCR2_0 subprocess.

    Chat and embeddings are small enough to stay resident together the rest of the time
    (see the VRAM profile in RAG_IMPLEMENTATION_PLAN.md) and normally run under
    `--models-max 2` so both stay loaded. Only a GPU-resident VLM like GOT-OCR2_0 needs
    both unloaded first; RapidOCR/tesseract run on CPU and never need this.
    """

    def __init__(self) -> None:
        self.events: list[dict[str, str]] = []
        self._unloaded: list[str] = []

    def _call(self, action: str, model: str) -> None:
        _router_request(action, model)
        self.events.append({"action": action, "model": model})

    def __enter__(self) -> "GotOcrVramSwap":
        try:
            for model in dict.fromkeys((LLM_MODEL, EMBEDDINGS_MODEL)):
                self._call("unload", model)
                self._unloaded.append(model)
        except Exception:
            self._restore()
            raise
        return self

    def _restore(self) -> list[str]:
        errors: list[str] = []
        for model in reversed(self._unloaded):
            if model == LLM_MODEL and not MODEL_ROUTER_RESTORE_CHAT:
                continue
            try:
                self._call("load", model)
            except Exception as error:
                errors.append(str(error))
        return errors

    def __exit__(self, exc_type, exc, traceback) -> bool:
        errors = self._restore()
        if errors and exc_type is None:
            raise RuntimeError("Failed to restore models after GOT-OCR2_0: " + "; ".join(errors))
        return False


def orchestrated_conversion(function):
    @wraps(function)
    def wrapper(*args, **kwargs):
        result = function(*args, **kwargs)
        if MODEL_ROUTER_ENABLED:
            store = kwargs.get("store") or (args[1] if len(args) > 1 else None) or RagStore()
            from backend.rag.embeddings import index_embeddings

            result["embedding_index"] = index_embeddings(store=store)
        return result

    return wrapper


def _got_ocr_image(path: Path) -> str:
    required = ("config.json", "model.safetensors")
    missing = [name for name in required if not (OCR_MODEL_DIR / name).exists()]
    if missing:
        raise RuntimeError(
            f"GOT-OCR2_0 weights not found in {OCR_MODEL_DIR} (missing {', '.join(missing)}). "
            "Run: DOWNLOAD_OCR_MODEL=1 ./download-rag-models.sh, then "
            "../venv/bin/pip install -r backend/requirements-ocr-got.txt"
        )
    worker = Path(__file__).with_name("got_ocr_worker.py")

    def _run() -> subprocess.CompletedProcess:
        return subprocess.run(
            [sys.executable, str(worker), str(path), str(OCR_MODEL_DIR)],
            capture_output=True,
            text=True,
            timeout=600,
        )

    if MODEL_ROUTER_ENABLED:
        with GotOcrVramSwap():
            result = _run()
    else:
        result = _run()
    if result.returncode != 0:
        raise RuntimeError(f"GOT-OCR2_0 worker failed: {result.stderr.strip()[-2000:]}")
    return result.stdout.strip()


def _ocr_image(path: Path, provider: str = OCR_PROVIDER) -> str:
    if not OCR_ENABLED:
        raise RuntimeError("OCR is disabled by OCR_ENABLED")
    data = path.read_bytes()
    cache = RAG_CACHE_DIR / "ocr" / f"{_cache_key(data, provider)}.md"
    if cache.exists():
        metadata_path = cache.with_suffix(".json")
        if not metadata_path.exists():
            metadata_path.write_text(json.dumps({"provider": provider, "boxes": [], "legacy_cache": True}), encoding="utf-8")
        return cache.read_text(encoding="utf-8")
    cache.parent.mkdir(parents=True, exist_ok=True)
    if provider == "rapidocr":
        from rapidocr_onnxruntime import RapidOCR

        engine = RapidOCR()
        result, _ = engine(str(path))
        text = "\n".join(str(line[1]) for line in (result or []))
        boxes = [{"box": line[0], "text": line[1], "confidence": line[2]} for line in (result or [])]
    elif provider == "tesseract":
        executable = shutil.which("tesseract")
        if not executable:
            raise RuntimeError("tesseract is required for OCR_PROVIDER=tesseract")
        result = subprocess.run([executable, str(path), "stdout", "-l", "eng+spa"], check=True, capture_output=True, text=True, timeout=300)
        text = result.stdout.strip()
        boxes = []
    elif provider == "got_ocr":
        text = _got_ocr_image(path)
        boxes = []
    else:
        raise RuntimeError(f"Unsupported OCR provider: {provider}; use rapidocr/tesseract/got_ocr")
    cache.write_text(text + "\n", encoding="utf-8")
    cache.with_suffix(".json").write_text(json.dumps({"provider": provider, "boxes": boxes}, ensure_ascii=False), encoding="utf-8")
    return text


def _pdf_to_markdown(path: Path) -> tuple[str, dict[str, Any]]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages: list[str] = []
    ocr_pages: list[int] = []
    ocr_artifacts: list[str] = []
    for index, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        printable_ratio = sum(char.isprintable() for char in text) / max(1, len(text))
        if len(text) < 10 or printable_ratio < 0.85:
            ocr_pages.append(index)
            renderer = shutil.which("pdftoppm")
            if not renderer:
                raise RuntimeError("pdftoppm is required to OCR scanned PDF pages")
            with tempfile.TemporaryDirectory() as temp:
                prefix = Path(temp) / "page"
                subprocess.run([renderer, "-f", str(index), "-singlefile", "-png", "-r", "200", str(path), str(prefix)], check=True, timeout=300)
                image_path = prefix.with_suffix(".png")
                text = _ocr_image(image_path)
                key = _cache_key(image_path.read_bytes(), OCR_PROVIDER)
                ocr_artifacts.append((RAG_CACHE_DIR / "ocr" / f"{key}.json").relative_to(WORKSPACE_ROOT).as_posix())
        pages.append(f"## Pagina {index}\n\n{text}")
    return f"# {path.stem}\n\n" + "\n\n".join(pages) + "\n", {"pages": len(pages), "ocr_pages": ocr_pages, "ocr_artifacts": ocr_artifacts}


def _docx_to_markdown(path: Path) -> tuple[str, dict[str, Any]]:
    from docx import Document

    document = Document(path)
    output = [f"# {path.stem}", ""]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name.lower() if paragraph.style else ""
        if style.startswith("heading"):
            level = min(6, int(style.rsplit(" ", 1)[-1]) if style.rsplit(" ", 1)[-1].isdigit() else 2)
            output.append("#" * level + " " + text)
        elif "list" in style:
            output.append("- " + text)
        else:
            output.append(text)
        output.append("")
    for table in document.tables:
        rows = [[cell.text.strip().replace("|", "\\|") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        output.append("| " + " | ".join(rows[0]) + " |")
        output.append("| " + " | ".join("---" for _ in rows[0]) + " |")
        output.extend("| " + " | ".join(row) + " |" for row in rows[1:])
        output.append("")
    return "\n".join(output), {"paragraphs": len(document.paragraphs), "tables": len(document.tables)}



@orchestrated_conversion
def convert_to_markdown(path: Path, store: RagStore | None = None) -> dict[str, Any]:
    path = _safe_input(path)
    store = store or RagStore()
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        markdown, metadata = _pdf_to_markdown(path)
        converter = f"pypdf+{OCR_PROVIDER}" if metadata["ocr_pages"] else "pypdf"
    elif suffix == ".docx":
        markdown, metadata = _docx_to_markdown(path)
        converter = "python-docx"
    elif suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"}:
        markdown = f"# {path.stem}\n\n{_ocr_image(path)}\n"
        key = _cache_key(path.read_bytes(), OCR_PROVIDER)
        metadata = {"pages": 1, "ocr_pages": [1], "ocr_artifacts": [(RAG_CACHE_DIR / "ocr" / f"{key}.json").relative_to(WORKSPACE_ROOT).as_posix()]}
        converter = OCR_PROVIDER
    else:
        raise ValueError(f"Unsupported document type: {suffix}")
    relative = path.relative_to(WORKSPACE_ROOT)
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    output = RAG_SOURCES_DIR / relative.parent / f"{path.stem}-{digest[:12]}.md"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(markdown, encoding="utf-8")
    source_id = hashlib.sha256(str(relative).encode()).hexdigest()
    with store.connect() as db:
        db.execute(
            """INSERT OR REPLACE INTO sources(id,original_path,normalized_markdown_path,converter,converter_version,ocr_model,metadata_json,created_at)
               VALUES(?,?,?,?,?,?,?,?)""",
            (source_id, relative.as_posix(), output.relative_to(WORKSPACE_ROOT).as_posix(), converter, "1", OCR_MODEL if metadata.get("ocr_pages") else None, json.dumps(metadata), time.time()),
        )
    from backend.rag.ingest import ingest_documents

    ingestion = ingest_documents(RAG_SOURCES_DIR, store)
    return {"source": relative.as_posix(), "markdown": output.relative_to(WORKSPACE_ROOT).as_posix(), "converter": converter, "ingestion": ingestion, **metadata}
